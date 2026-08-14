import requests
import os
import re
import json
import zipfile
import tempfile
import pypdf
import pdfplumber
import docx
from pathlib import Path

BASE_DIR = Path(r"C:\Users\11429149760\.gemini\antigravity\scratch\salesforce_tr_crawler")
DOWNLOADS_DIR = BASE_DIR / "downloads"
DOWNLOADS_DIR.mkdir(parents=True, exist_ok=True)

def get_process_files(session, item):
    cnpj = item.get('orgao_cnpj')
    ano = item.get('ano')
    seq = item.get('numero_sequencial')
    
    if not (cnpj and ano and seq) and item.get('item_url'):
        parts = item['item_url'].strip('/').split('/')
        if len(parts) >= 4 and parts[0] == 'compras':
            cnpj, ano, seq = parts[1], parts[2], parts[3]
            
    if not (cnpj and ano and seq):
        return []
        
    url = f"https://pncp.gov.br/api/pncp/v1/orgaos/{cnpj}/compras/{ano}/{seq}/arquivos"
    try:
        res = session.get(url, timeout=15)
        if res.status_code == 200:
            return res.json()
        else:
            return []
    except Exception as e:
        print(f"[FILES ERROR] Failed to list files for {cnpj}/{ano}/{seq}: {e}")
        return []

def download_file(session, file_url, save_path):
    try:
        res = session.get(file_url, stream=True, timeout=30)
        if res.status_code == 200:
            with open(save_path, 'wb') as f:
                for chunk in res.iter_content(chunk_size=8192):
                    if chunk:
                        f.write(chunk)
            return True
        else:
            print(f"[DOWNLOAD FAIL] Status {res.status_code} for {file_url}")
            return False
    except Exception as e:
        print(f"[DOWNLOAD EXCEPTION] {file_url}: {e}")
        return False

def extract_text_from_file(file_path):
    if not os.path.exists(file_path) or os.path.getsize(file_path) == 0:
        return ""
        
    # Check magic bytes
    with open(file_path, 'rb') as f:
        header = f.read(10)
        
    is_zip_header = header.startswith(b'PK\x03\x04')
    is_pdf_header = header.startswith(b'%PDF')
    ext = os.path.splitext(file_path)[1].lower()
    text = ""
    
    # If ZIP file (could be DOCX or actual ZIP archive)
    if is_zip_header:
        try:
            with zipfile.ZipFile(file_path, 'r') as z:
                names = z.namelist()
                # Check if DOCX
                if 'word/document.xml' in names:
                    try:
                        doc = docx.Document(file_path)
                        return "\n".join([p.text for p in doc.paragraphs if p.text])
                    except Exception:
                        pass
                # Otherwise extract inner files
                extracted_texts = []
                with tempfile.TemporaryDirectory() as tmpdir:
                    z.extractall(tmpdir)
                    for root, _, files in os.walk(tmpdir):
                        for file in files:
                            inner_path = os.path.join(root, file)
                            t = extract_text_from_file(inner_path)
                            if t.strip():
                                extracted_texts.append(f"--- [ANEXO COMPACTADO: {file}] ---\n" + t)
                if extracted_texts:
                    return "\n\n".join(extracted_texts)
        except Exception as e:
            print(f"[ZIP READ FAIL] {file_path}: {e}")
            
    # PDF Processing
    if is_pdf_header or ext == '.pdf':
        try:
            reader = pypdf.PdfReader(file_path)
            pages_text = []
            for page in reader.pages:
                t = page.extract_text()
                if t:
                    pages_text.append(t)
            text = "\n".join(pages_text)
        except Exception as e:
            text = ""
            
        if len(text.strip()) < 50:
            try:
                with pdfplumber.open(file_path) as pdf:
                    pages_text = []
                    for page in pdf.pages:
                        t = page.extract_text()
                        if t:
                            pages_text.append(t)
                    text = "\n".join(pages_text)
            except Exception as e:
                pass
                
    elif ext in ['.docx', '.doc']:
        try:
            doc = docx.Document(file_path)
            text = "\n".join([p.text for p in doc.paragraphs if p.text])
        except Exception as e:
            pass
            
    elif ext in ['.txt', '.html', '.xml']:
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                text = f.read()
        except Exception as e:
            pass
            
    return text

def process_item_downloads_and_extraction(session, item):
    item_id = item.get('id') or item.get('numero_controle_pncp') or "proc_unknown"
    proc_dir = DOWNLOADS_DIR / item_id
    proc_dir.mkdir(parents=True, exist_ok=True)
    
    files_info = get_process_files(session, item)
    downloaded_files = []
    combined_texts = []
    
    print(f"\n[DOWNLOAD] Processo {item_id} ({item.get('title')}) -> {len(files_info)} arquivos no PNCP")
    
    for f in files_info:
        file_url = f.get('url')
        file_title = f.get('titulo') or "documento"
        file_type = f.get('tipoDocumentoNome') or "Edital"
        
        safe_title = re.sub(r'[^\w\.-]', '_', file_title)
        if not safe_title.endswith(('.pdf', '.docx', '.doc', '.zip')):
            safe_title += '.pdf'
            
        save_path = proc_dir / safe_title
        
        if file_url:
            print(f" -> Baixando {safe_title}...")
            success = download_file(session, file_url, save_path)
            if success:
                extracted_text = extract_text_from_file(save_path)
                downloaded_files.append({
                    'title': file_title,
                    'type': file_type,
                    'file_name': safe_title,
                    'local_path': str(save_path),
                    'url': file_url,
                    'text_length': len(extracted_text),
                    'sample_text': extracted_text[:500]
                })
                if extracted_text:
                    combined_texts.append(f"--- ARQUIVO: {file_title} ({file_type}) ---\n" + extracted_text)
                    
    desc = item.get('description') or ""
    if desc:
        combined_texts.insert(0, f"--- DESCRIÇÃO DO OBJETO NO PNCP ---\n{desc}")
        
    full_extracted_text = "\n\n".join(combined_texts)
    
    return {
        'item_id': item_id,
        'downloaded_files': downloaded_files,
        'full_extracted_text': full_extracted_text,
        'has_files': len(downloaded_files) > 0
    }
